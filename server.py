#!/usr/bin/env python3
"""
暖暖小窝 - 局域网协作服务器
提供 HTTP 文件服务和 WebSocket 房间管理

启动方式: python3 server.py
默认端口: HTTP 9000, WebSocket 9001
"""

import asyncio
import json
import socket
import threading
import http.server
import os
import sys
import io
from datetime import datetime
from urllib.parse import urlparse, parse_qs

try:
    import websockets
except ImportError:
    print("错误: 需要 websockets 库")
    print("安装: pip3 install websockets")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("提示: python-docx 未安装，Word 导出不可用")

# 配置
HTTP_PORT = 9000
WS_PORT = 9001
HOST = '0.0.0.0'
FIXED_ROOM_ID = '暖暖小窝'
ALLOWED_NAMES = ['大灰狼', '懒洋洋']
EXTENSION_VERSION = '1.1.1'
HISTORY_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'chat_history.json')

# ============================================================
# HTTP 静态文件服务器 + API
# ============================================================
class RequestHandler(http.server.SimpleHTTPRequestHandler):
    def log_message(self, format, *args):
        pass

    def do_GET(self):
        parsed = urlparse(self.path)

        if parsed.path == '/api/export':
            self.handle_export()
        elif parsed.path == '/api/history':
            self.handle_history_json()
        elif parsed.path == '/api/version':
            self.handle_version()
        elif parsed.path == '/extension/' or parsed.path == '/extension':
            self.handle_extension_page()
        elif parsed.path == '/extension/extension.zip':
            self.handle_extension_download()
        else:
            super().do_GET()

    def handle_version(self):
        self.send_response(200)
        self.send_header('Content-Type', 'application/json; charset=utf-8')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(json.dumps({
            'version': EXTENSION_VERSION,
            'update_url': '/extension/'
        }).encode('utf-8'))

    def handle_extension_page(self):
        # 插件更新页面
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>暖暖小窝 - 插件更新</title>
  <style>
    * {{ margin:0; padding:0; box-sizing:border-box; }}
    body {{ font-family:'Nunito',-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif; background:#F8F8F0; color:#794F27; display:flex; align-items:center; justify-content:center; min-height:100vh; }}
    .card {{ background:#F7F3DF; border-radius:24px; padding:40px; max-width:480px; width:90%; text-align:center; box-shadow:0 8px 32px rgba(0,0,0,0.1); }}
    .icon {{ width:64px; height:64px; margin:0 auto 20px; background:#19C8B9; border-radius:50%; display:flex; align-items:center; justify-content:center; }}
    .icon svg {{ width:32px; height:32px; stroke:white; fill:none; stroke-width:2; }}
    h1 {{ font-size:24px; margin-bottom:8px; font-weight:700; }}
    .version {{ color:#19C8B9; font-size:14px; font-weight:700; margin-bottom:16px; }}
    p {{ color:#725D42; font-size:14px; line-height:1.6; margin-bottom:24px; }}
    .btn {{ display:inline-block; padding:14px 32px; background:#19C8B9; color:white; text-decoration:none; border-radius:50px; font-weight:700; font-size:16px; box-shadow:0 5px 0 0 #0EA89B; transition:all 0.15s; }}
    .btn:hover {{ transform:translateY(-2px); box-shadow:0 7px 0 0 #0EA89B; }}
    .btn:active {{ transform:translateY(2px); box-shadow:0 3px 0 0 #0EA89B; }}
    .steps {{ text-align:left; margin-top:24px; padding:16px; background:rgba(0,0,0,0.03); border-radius:16px; }}
    .steps h3 {{ font-size:14px; margin-bottom:12px; color:#794F27; }}
    .steps ol {{ padding-left:20px; font-size:13px; color:#725D42; line-height:1.8; }}
  </style>
</head>
<body>
  <div class="card">
    <div class="icon">
      <svg viewBox="0 0 24 24" stroke-linecap="round" stroke-linejoin="round"><path d="M3 9l9-7 9 7v11a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2z"/><polyline points="9 22 9 12 15 12 15 22"/></svg>
    </div>
    <h1>暖暖小窝</h1>
    <div class="version">最新版本 v{EXTENSION_VERSION}</div>
    <p>下载最新的浏览器插件安装包，解压后在浏览器扩展页面重新加载即可。</p>
    <a class="btn" href="/extension/extension.zip" download>下载插件</a>
    <div class="steps">
      <h3>安装步骤</h3>
      <ol>
        <li>下载并解压 ZIP 文件</li>
        <li>打开 <code>chrome://extensions/</code>（夸克浏览器：<code>quark://extensions/</code>）</li>
        <li>开启右上角「开发者模式」</li>
        <li>点击「加载已解压的扩展程序」</li>
        <li>选择解压后的文件夹</li>
      </ol>
    </div>
  </div>
</body>
</html>"""
        self.send_response(200)
        self.send_header('Content-Type', 'text/html; charset=utf-8')
        self.end_headers()
        self.wfile.write(html.encode('utf-8'))

    def handle_extension_download(self):
        # 从 extension 目录读取 ZIP 文件
        zip_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'extension.zip')
        if os.path.exists(zip_path):
            self.send_response(200)
            self.send_header('Content-Type', 'application/zip')
            self.send_header('Content-Disposition', 'attachment; filename="extension.zip"')
            self.send_header('Content-Length', str(os.path.getsize(zip_path)))
            self.end_headers()
            with open(zip_path, 'rb') as f:
                self.wfile.write(f.read())
        else:
            self.send_response(404)
            self.send_header('Content-Type', 'text/plain')
            self.end_headers()
            self.wfile.write(b'extension.zip not found')

    def handle_history_json(self):
        self.send_response(200)
        self.send_header('Content-Type', 'application/json; charset=utf-8')
        self.end_headers()
        self.wfile.write(json.dumps(room_state['chat_history'], ensure_ascii=False, indent=2).encode('utf-8'))

    def handle_export(self):
        if not HAS_DOCX:
            self.send_response(500)
            self.send_header('Content-Type', 'text/plain; charset=utf-8')
            self.end_headers()
            self.wfile.write('python-docx 未安装'.encode('utf-8'))
            return

        doc = Document()

        # 标题
        title = doc.add_heading('暖暖小窝 - 聊天记录', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 导出时间
        p = doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_paragraph('')

        chat_history = room_state.get('chat_history', {})

        if not chat_history:
            doc.add_paragraph('暂无聊天记录')
        else:
            for chat_key, messages in sorted(chat_history.items()):
                # 聊天双方
                parts = chat_key.split('<->')
                if len(parts) == 2:
                    doc.add_heading(f'{parts[0]} <-> {parts[1]}', level=2)
                else:
                    doc.add_heading(chat_key, level=2)

                for msg in messages:
                    sender = msg.get('from', '未知')
                    text = msg.get('text', '')
                    time_str = msg.get('time', '')
                    try:
                        dt = datetime.fromisoformat(time_str)
                        time_display = dt.strftime('%m-%d %H:%M')
                    except:
                        time_display = time_str

                    p = doc.add_paragraph()
                    run_time = p.add_run(f'[{time_display}] ')
                    run_time.font.size = Pt(9)
                    run_time.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

                    run_name = p.add_run(f'{sender}: ')
                    run_name.bold = True
                    run_name.font.size = Pt(11)

                    run_text = p.add_run(text)
                    run_text.font.size = Pt(11)

                doc.add_paragraph('')

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        self.send_response(200)
        self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.send_header('Content-Disposition', f'attachment; filename="chat_history_{datetime.now().strftime("%Y%m%d")}.docx"')
        self.end_headers()
        self.wfile.write(buffer.read())


def run_http():
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    server = http.server.HTTPServer((HOST, HTTP_PORT), RequestHandler)
    server.serve_forever()

# ============================================================
# WebSocket 房间管理
# ============================================================
room_state = {
    'chat_history': {},
    'todos': [],
    'foods': ['黄焖鸡', '麻辣烫', '牛肉面', '寿司', '炸鸡', '沙拉', '饺子', '螺蛳粉']
}
members = {}    # username -> websocket
clients = {}    # websocket -> username


def save_history():
    """保存聊天记录到文件"""
    try:
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(room_state['chat_history'], f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f'[保存历史失败] {e}')


def load_history():
    """从文件加载聊天记录"""
    global room_state
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                room_state['chat_history'] = json.load(f)
            print(f'[加载历史] {len(room_state["chat_history"])} 条对话')
        except Exception as e:
            print(f'[加载历史失败] {e}')


async def handle_client(websocket):
    username = None

    try:
        async for message in websocket:
            try:
                data = json.loads(message)
                msg_type = data.get('type')

                if msg_type == 'join':
                    username = await handle_join(websocket, data)
                elif username:
                    await handle_message(websocket, username, data)

            except json.JSONDecodeError:
                pass
            except Exception as e:
                print(f'[WS 错误] {e}')

    except websockets.exceptions.ConnectionClosed:
        pass
    finally:
        if username:
            await handle_leave(websocket, username)


async def handle_join(websocket, data):
    username = data.get('username', '').strip()

    if not username:
        await websocket.send(json.dumps({'type': 'error', 'message': '请输入昵称'}))
        return None

    # 昵称限制
    if username not in ALLOWED_NAMES:
        await websocket.send(json.dumps({'type': 'error', 'message': '开发中'}))
        return None

    # 如果昵称已在线，踢掉旧连接
    if username in members:
        old_ws = members[username]
        try:
            await old_ws.send(json.dumps({'type': 'kicked', 'message': '相同昵称从其他设备登录'}))
            await old_ws.close()
        except:
            pass
        if old_ws in clients:
            del clients[old_ws]

    members[username] = websocket
    clients[websocket] = username

    await websocket.send(json.dumps({
        'type': 'room-joined',
        'room_id': FIXED_ROOM_ID,
        'username': username,
        'members': list(members.keys()),
        'state': room_state,
        'lan_ips': get_lan_ips()
    }))

    # 通知其他人（包含在线状态）
    await broadcast({
        'type': 'member-joined',
        'username': username,
        'members': list(members.keys())
    }, exclude=websocket)

    print(f'[加入] {username} (共 {len(members)} 人在线)')
    return username


async def handle_message(websocket, username, data):
    msg_type = data.get('type')

    if msg_type == 'chat':
        msg_id = data.get('msg_id', f"msg-{datetime.now().timestamp()}")
        text = data.get('text', '')
        to_user = data.get('to_user')

        chat_key = get_chat_key(username, to_user) if to_user else f"{username}->all"
        if chat_key not in room_state['chat_history']:
            room_state['chat_history'][chat_key] = []

        msg = {
            'id': msg_id,
            'text': text,
            'from': username,
            'to': to_user,
            'time': datetime.now().isoformat(),
            'read': False
        }

        # 端到端加密字段透传
        if data.get('encrypted'):
            msg['encrypted'] = True
            msg['ciphertext'] = data.get('ciphertext', '')
            msg['iv'] = data.get('iv', '')
            msg['text'] = '🔒 加密消息'  # 服务端只看到占位文本

        room_state['chat_history'][chat_key].append(msg)
        save_history()

        if to_user:
            target_ws = members.get(to_user)
            if target_ws:
                await target_ws.send(json.dumps({'type': 'chat', 'message': msg}))
        else:
            await broadcast({'type': 'chat', 'message': msg}, exclude=websocket)

        await websocket.send(json.dumps({'type': 'chat-receipt', 'msg_id': msg_id}))

    elif msg_type == 'e2ee-key-offer':
        # 密钥交换请求：A -> Server -> B
        to_user = data.get('to_user')
        if to_user:
            target_ws = members.get(to_user)
            if target_ws:
                await target_ws.send(json.dumps({
                    'type': 'e2ee-key-offer',
                    'from': username,
                    'public_key': data.get('public_key', '')
                }))

    elif msg_type == 'e2ee-key-answer':
        # 密钥交换回复：B -> Server -> A
        to_user = data.get('to_user')
        if to_user:
            target_ws = members.get(to_user)
            if target_ws:
                await target_ws.send(json.dumps({
                    'type': 'e2ee-key-answer',
                    'from': username,
                    'public_key': data.get('public_key', '')
                }))

    elif msg_type == 'chat-read':
        # 已读回执
        chat_key = data.get('chat_key', '')
        read_msgs = data.get('msg_ids', [])

        if chat_key in room_state['chat_history']:
            for msg in room_state['chat_history'][chat_key]:
                if msg['id'] in read_msgs:
                    msg['read'] = True
            save_history()

        # 通知发送者消息已读
        to_user = data.get('to_user', '')
        target_ws = members.get(to_user)
        if target_ws:
            await target_ws.send(json.dumps({
                'type': 'chat-read-notify',
                'chat_key': chat_key,
                'msg_ids': read_msgs,
                'read_by': username
            }))

    elif msg_type == 'todo-add':
        todo = data.get('todo', {})
        todo['created_by'] = username
        room_state['todos'].append(todo)
        await broadcast({'type': 'todo-added', 'todo': todo}, exclude=websocket)

    elif msg_type == 'todo-toggle':
        todo_id = data.get('id')
        for t in room_state['todos']:
            if t['id'] == todo_id:
                t['done'] = not t.get('done', False)
                break
        await broadcast({'type': 'todo-toggled', 'id': todo_id}, exclude=websocket)

    elif msg_type == 'todo-delete':
        todo_id = data.get('id')
        room_state['todos'] = [t for t in room_state['todos'] if t['id'] != todo_id]
        await broadcast({'type': 'todo-deleted', 'id': todo_id}, exclude=websocket)

    elif msg_type == 'food-add':
        food = data.get('food', '')
        if food and food not in room_state['foods']:
            room_state['foods'].append(food)
        await broadcast({'type': 'food-added', 'food': food}, exclude=websocket)

    elif msg_type == 'food-remove':
        food = data.get('food', '')
        if food in room_state['foods']:
            room_state['foods'].remove(food)
        await broadcast({'type': 'food-removed', 'food': food}, exclude=websocket)

    elif msg_type == 'sync-request':
        await websocket.send(json.dumps({
            'type': 'sync',
            'state': room_state,
            'members': list(members.keys())
        }))

    else:
        await broadcast(data, exclude=websocket)


async def handle_leave(websocket, username):
    if username in members:
        del members[username]

    if websocket in clients:
        del clients[websocket]

    await broadcast({
        'type': 'member-left',
        'username': username,
        'members': list(members.keys())
    })

    print(f'[离开] {username} (共 {len(members)} 人在线)')


async def broadcast(data, exclude=None):
    message = json.dumps(data)
    for ws in list(members.values()):
        if ws != exclude:
            try:
                await ws.send(message)
            except:
                pass


def get_chat_key(a, b):
    return '<->'.join(sorted([a, b]))


def get_lan_ips():
    ips = set()
    try:
        for info in socket.getaddrinfo(socket.gethostname(), None, socket.AF_INET):
            ip = info[4][0]
            if not ip.startswith('127.'):
                ips.add(ip)
    except:
        pass
    return list(ips)


async def ws_main():
    async with websockets.serve(handle_client, HOST, WS_PORT):
        print('')
        print('  ========================================')
        print('  🌿  暖暖小窝 - 局域网服务器已启动')
        print('  ========================================')
        print('')
        print(f'  📱 本机访问:  http://localhost:{HTTP_PORT}')
        print('')
        ips = get_lan_ips()
        if ips:
            print(f'  📡 局域网访问 (同一WiFi下的设备):')
            for ip in ips:
                print(f'     → http://{ip}:{HTTP_PORT}')
            print('')
        print(f'  🔌 WebSocket: ws://localhost:{WS_PORT}')
        print('')
        print('  📝 聊天记录保存到: chat_history.json')
        print(f'  📄 Word 导出: http://localhost:{HTTP_PORT}/api/export')
        print('')
        print('  按 Ctrl+C 停止服务器')
        print('')
        await asyncio.Future()


def main():
    load_history()

    t = threading.Thread(target=run_http, daemon=True)
    t.start()

    try:
        asyncio.run(ws_main())
    except KeyboardInterrupt:
        print('\n  服务器已停止\n')


if __name__ == '__main__':
    main()
